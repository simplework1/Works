from docx import Document
from docx.shared import Pt, RGBColor

# Function to copy the cell formatting
def copy_cell_formatting(source_cell, target_cell):
    target_cell.paragraphs[0].style = source_cell.paragraphs[0].style
    target_cell.paragraphs[0].alignment = source_cell.paragraphs[0].alignment
    if len(source_cell.paragraphs[0].runs) > 0:
        source_run = source_cell.paragraphs[0].runs[0]
        target_run = target_cell.paragraphs[0].runs[0]
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.size = source_run.font.size
        target_run.font.color.rgb = source_run.font.color.rgb
        target_run.font.name = source_run.font.name

# Function to add rows and columns to a table while maintaining formatting
def add_rows_and_columns(table, num_rows, num_cols):
    # Add the specified number of rows
    for _ in range(num_rows):
        new_row = table.add_row()
        for i in range(len(new_row.cells)):
            copy_cell_formatting(table.rows[0].cells[i], new_row.cells[i])
    
    # Add the specified number of columns
    for row in table.rows:
        for _ in range(num_cols):
            new_cell = row.add_cell()
            copy_cell_formatting(row.cells[-(num_cols + 1)], new_cell)

# Load the existing document to extract the table
doc_path = '/mnt/data/file-9NnwlpnnmO0xNXUMpcriEW3z'
doc = Document(doc_path)
existing_table = doc.tables[0]

# Add 2 rows and 1 column to the existing table while maintaining formatting
add_rows_and_columns(existing_table, 2, 1)

# Save the modified document
output_path = '/mnt/data/document_with_added_rows_columns.docx'
doc.save(output_path)

print(f"Document saved to {output_path}")
