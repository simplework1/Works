from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

# Function to copy cell formatting and shading
def copy_cell_formatting_and_shading(source_cell, target_cell):
    # Copy text formatting
    target_cell.paragraphs[0].style = source_cell.paragraphs[0].style
    target_cell.paragraphs[0].alignment = source_cell.paragraphs[0].alignment
    target_cell.vertical_alignment = source_cell.vertical_alignment

    # Clear existing runs in target cell and copy runs from source cell
    target_cell.paragraphs[0].clear()
    for source_paragraph in source_cell.paragraphs:
        for source_run in source_paragraph.runs:
            target_run = target_cell.paragraphs[0].add_run(source_run.text)
            target_run.bold = source_run.bold
            target_run.italic = source_run.italic
            target_run.underline = source_run.underline
            target_run.font.size = source_run.font.size
            if source_run.font.color and source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
            target_run.font.name = source_run.font.name

    # Copy cell shading
    source_shading = source_cell._element.xpath('.//w:shd')
    if source_shading:
        target_shading = OxmlElement('w:shd')
        for attr in source_shading[0].attrib:
            target_shading.set(attr, source_shading[0].get(attr))
        tcPr = target_cell._element.get_or_add_tcPr()
        tcPr.append(target_shading)

# Function to duplicate a table within the same document
def duplicate_table_within_document(source_table, document):
    # Create a new table with the same number of rows and columns
    new_table = document.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    new_table.style = source_table.style
    
    for row_idx, row in enumerate(source_table.rows):
        for col_idx, cell in enumerate(row.cells):
            new_cell = new_table.cell(row_idx, col_idx)
            copy_cell_formatting_and_shading(cell, new_cell)
            new_cell.text = cell.text
    
    return new_table

# Load the existing document to extract the table
source_doc_path = '/mnt/data/file-WwrlxJTY82VWd14mD5ZRK0Pp'
doc = Document(source_doc_path)
source_table = doc.tables[0]

# Duplicate the table within the same document
duplicate_table_within_document(source_table, doc)

# Save the modified document
output_path = '/mnt/data/document_with_duplicated_table.docx'
doc.save(output_path)

print(f"Document saved to {output_path}")
