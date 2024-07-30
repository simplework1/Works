from docx import Document

# Load the document
doc_path = '/mnt/data/file-rU7Y1IamFEmkE1jRQgIwFeWe'
doc = Document(doc_path)

# Function to add a footer
def add_footer(doc, text_lines):
    section = doc.sections[0]
    footer = section.footer
    for text in text_lines:
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(8)  # Set font size to match your requirement
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

# Text lines to be added to the footer
footer_text = [
    "Source: Partnership Deed of Viva Clothing dated 10 May 2019 provided by the Client",
    "PAN is a unique alphanumeric combination issued to all juristic entities identifiable under the Income Tax Act, 1961 by the Income Tax Department of India",
    "Source: Partnership Deed of Viva Clothing dated 10 May 2019 provided by the Client"
]

# Add footer to the document
add_footer(doc, footer_text)

# Save the document with changes
output_path = '/mnt/data/document_with_footer.docx'
doc.save(output_path)

print(f"Document saved to {output_path}")
